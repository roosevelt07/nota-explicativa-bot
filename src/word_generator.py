# src/word_generator.py
"""
Geração de arquivo Word (.docx) do relatório, usando o mesmo dicionário 'dados'
do PDF (montado em core.montar_dados_relatorio).

Objetivo:
- Replicar o conteúdo do PDF (seções, textos e tabelas);
- Gerar um arquivo de 1–2 páginas, sem exagero de quebras;
- Incluir o papel timbrado como imagem no topo do documento.
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import Dict, Any

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.utils import formatar_total_previdencia


def _configurar_estilo_normal(doc: Document) -> None:
    """Configura o estilo padrão do documento."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _add_heading(doc: Document, text: str) -> None:
    """Adiciona um título de seção, equivalente aos headings do PDF."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_paragrafo(doc: Document, text: str) -> None:
    """Parágrafo padrão."""
    if not text:
        return
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_table(doc: Document, headers, rows):
    """Tabela simples com cabeçalho em negrito."""
    if not rows:
        return

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    # Cabeçalho
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h

    for cell in hdr_cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Linhas de dados
    for row in rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
        for cell in row_cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    # Espaço após a tabela
    doc.add_paragraph("")


def _fmt_moeda_word(valor):
    """Formata float para R$ X.XXX,XX"""
    if not valor:
        return "-"
    try:
        val_float = float(valor)
        return f"R$ {val_float:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except (ValueError, TypeError):
        return str(valor)


def _add_papel_timbrado_topo(doc: Document) -> None:
    """
    Adiciona o papel timbrado como imagem no topo do documento (antes do texto).
    Usa o mesmo arquivo do PDF: assets/nota_explicativa_em_branco.png
    """
    base_dir = Path(__file__).resolve().parent.parent
    template_path = base_dir / "assets" / "nota_explicativa_em_branco.png"

    if not template_path.exists():
        return

    # Usa a largura útil da página (largura menos margens) para a imagem
    section = doc.sections[0]
    largura_util = section.page_width - section.left_margin - section.right_margin

    # add_picture aceita o valor em EMU (page_width já está em EMU)
    doc.add_picture(str(template_path), width=largura_util)
    # Pequeno espaço depois da imagem
    doc.add_paragraph("")


def gerar_docx_bytes(dados: Dict[str, Any]) -> bytes:
    """
    Gera o arquivo .docx em memória a partir do dicionário 'dados'.
    Retorna os bytes do arquivo para uso no st.download_button.
    """
    doc = Document()
    _configurar_estilo_normal(doc)

    # Papel timbrado no topo
    _add_papel_timbrado_topo(doc)

    # ======================= CABEÇALHO / DADOS BÁSICOS =======================

    _add_paragrafo(doc, f"Data do relatório: {dados['data_relatorio']}")
    _add_paragrafo(doc, f"Requerente: {dados['requerente']}")
    _add_paragrafo(doc, f"CNPJ: {dados['cnpj']}")
    _add_paragrafo(doc, f"Tributação: {dados['tributacao']}")
    _add_paragrafo(doc, f"Certificado Digital: {dados['certificado_digital']}")
    doc.add_paragraph("")

    intro = (
        "Este relatório tem como objetivo acompanhar os débitos pendentes relacionados à entidade "
        "empresarial destacada acima, destacando os principais pontos sobre a situação fiscal, os "
        "valores devidos, datas de vencimento e providências necessárias para regularização. Nos "
        "casos em que haja desacordo com os débitos e irregularidades apresentadas ou já tenha sido "
        "efetuado o pagamento, favor entrar em contato conosco para a resolução da pendência."
    )
    _add_paragrafo(doc, intro)
    doc.add_paragraph("")

    _add_heading(doc, "DÉBITOS IDENTIFICADOS")
    _add_paragrafo(
        doc,
        "Abaixo, estão listados os débitos pendentes e a situação atual da empresa:",
    )
    doc.add_paragraph("")

    # ========================= RECEITA FEDERAL =========================

    _add_heading(doc, "RECEITA FEDERAL")
    
    # Total de Previdência (OBJETIVO 3) - SOMENTE o total, sem tabela completa
    texto_total_previdencia = formatar_total_previdencia(dados)
    _add_paragrafo(doc, texto_total_previdencia)
    doc.add_paragraph("")
    
    # PGFN Previdência (OBJETIVO 1)
    if "receita_federal" in dados and dados["receita_federal"]:
        receita = dados["receita_federal"]
        pgfn_previdencia = receita.get("pgfn_previdencia", {})
        if pgfn_previdencia.get("existe"):
            receitas_list = pgfn_previdencia.get("receitas", [])
            receitas_str = "; ".join(receitas_list) if receitas_list else "Não identificado"
            
            _add_paragrafo(doc, "PGFN Previdência")
            _add_paragrafo(doc, f"Receita: {receitas_str}")
            
            info_adicional = pgfn_previdencia.get("informacoes_adicionais_usuario", "")
            if info_adicional and info_adicional.strip():
                _add_paragrafo(doc, "Informações adicionais:")
                # Divide em parágrafos se houver quebras de linha
                for linha in info_adicional.split('\n'):
                    if linha.strip():
                        _add_paragrafo(doc, linha.strip())
            else:
                # Se não houver informações, mostra "(não informado)"
                _add_paragrafo(doc, "Informações adicionais: (não informado)")
            
            doc.add_paragraph("")
    
    _add_paragrafo(doc, dados.get("bloco_receita_federal", ""))
    _add_paragrafo(doc, f"Data da consulta: {dados['data_consulta_rf']}")
    doc.add_paragraph("")

    # ============================= SEFAZ ==============================

    _add_heading(doc, "SEFAZ")
    
    # Consolida linhas manuais com linhas extraídas automaticamente
    linhas_finais_sefaz = []
    
    # A) Adiciona linhas manuais (se houver)
    sefaz_rows = dados.get("sefaz_rows") or []
    if sefaz_rows:
        linhas_finais_sefaz.extend(sefaz_rows)
    
    # B) Adiciona dados estruturados do Parser SEFAZ
    if "sefaz_estadual" in dados and dados["sefaz_estadual"]:
        sefaz = dados["sefaz_estadual"]
        pendencias = sefaz.get("pendencias_identificadas", {})
        
        # IPVA
        for item in pendencias.get("ipva", []):
            desc = f"IPVA {item.get('exercicio', '')}"
            ref = item.get('placa', '')
            val = _fmt_moeda_word(item.get('valor_total', 0))
            linhas_finais_sefaz.append([desc, ref, val])
        
        # Fronteira/Antecipado
        for item in pendencias.get("icms_fronteira_antecipado", []):
            desc = item.get('descricao', 'ICMS Antecipado')
            ref = item.get('periodo_referencia', '')
            val = _fmt_moeda_word(item.get('valor_total', 0))
            linhas_finais_sefaz.append([desc, ref, val])
        
        # Débitos Fiscais
        for item in pendencias.get("debitos_fiscais_autuacoes", []):
            desc = f"Autuação {item.get('natureza_debito', '')}"
            ref = item.get('periodo', '')
            val = _fmt_moeda_word(item.get('valor_consolidado', 0))
            linhas_finais_sefaz.append([desc, ref, val])
    
    # Renderiza Tabela ou Mensagem "Sem Débitos"
    if linhas_finais_sefaz:
        _add_table(
            doc,
            ["Descrição do Débito / Pendência", "Período / Placa", "Valor / Situação"],
            linhas_finais_sefaz,
        )
    else:
        # Verifica se o parser identificou explicitamente como Regular
        status_geral = dados.get("sefaz_estadual", {}).get("cabecalho_documento", {}).get("situacao_geral", "")
        if "REGULAR" in status_geral.upper():
            _add_paragrafo(doc, "✅ Situação REGULAR (Certidão Negativa Emitida).")
        else:
            _add_paragrafo(doc, "Sem débitos informados ou identificados.")
    
    # Itens adicionais (manual)
    manual_sefaz = dados.get("sefaz", {}).get("itens_adicionais_manuais", "").strip()
    if manual_sefaz:
        _add_paragrafo(doc, "Itens adicionais (manual):")
        for linha in manual_sefaz.split("\n"):
            if linha.strip():
                _add_paragrafo(doc, linha.strip())
    else:
        _add_paragrafo(doc, "Itens adicionais (manual): (não informado)")
    
    _add_paragrafo(doc, f"Data da consulta: {dados['data_consulta_sefaz']}")
    doc.add_paragraph("")

    # ======================= DÉBITOS MUNICIPAIS =======================

    _add_heading(doc, "DÉBITOS MUNICIPAIS")
    municipais_rows = dados.get("municipais_rows") or []
    if municipais_rows:
        _add_table(
            doc,
            ["Descrição do Débito", "Período", "Valor", "Status"],
            municipais_rows,
        )
    else:
        _add_paragrafo(doc, "Sem débitos informados.")
    
    # Débitos municipais (manual)
    manual_mun = dados.get("debitos_municipais", {}).get("texto_manual", "").strip()
    if manual_mun:
        _add_paragrafo(doc, "Débitos municipais (manual):")
        for linha in manual_mun.split("\n"):
            if linha.strip():
                _add_paragrafo(doc, linha.strip())
    else:
        _add_paragrafo(doc, "Débitos municipais (manual): (não informado)")
    
    _add_paragrafo(doc, f"Data da consulta: {dados['data_consulta_municipal']}")
    doc.add_paragraph("")

    # =============================== FGTS =============================

    _add_heading(doc, "FGTS")
    
    # Dados estruturados do FGTS
    if "fgts" in dados and dados["fgts"]:
        fgts_data = dados["fgts"]
        crf = fgts_data.get("crf_detalhes", {})
        pendencias = fgts_data.get("pendencias_financeiras", {})
        
        # Resumo do Certificado
        if crf.get("numero_certificacao"):
            resumo_fgts_rows = [[
                crf.get("situacao_atual", "-"),
                f"{crf.get('validade_inicio','')} a {crf.get('validade_fim','')}",
                crf.get("numero_certificacao", "-")
            ]]
            _add_table(doc, ["Situação", "Validade", "Certificação"], resumo_fgts_rows)
            doc.add_paragraph("")
        
        # Tabela de Débitos do FGTS
        lista_debitos = pendencias.get("lista_debitos", [])
        if lista_debitos:
            tabela_fgts_rows = []
            for debito in lista_debitos:
                tabela_fgts_rows.append([
                    debito.get("competencia", "-"),
                    _fmt_moeda_word(debito.get("valor_estimado", 0)),
                    debito.get("situacao", "EM ABERTO")
                ])
            _add_table(doc, ["Competência", "Valor", "Situação"], tabela_fgts_rows)
            doc.add_paragraph("")
        elif crf.get("situacao_atual") == "REGULAR":
            _add_paragrafo(doc, "✅ Situação REGULAR - Não há débitos pendentes.")
            doc.add_paragraph("")
    
    _add_paragrafo(doc, dados.get("bloco_fgts", ""))
    _add_paragrafo(doc, f"Data da consulta: {dados['data_consulta_fgts']}")
    doc.add_paragraph("")

    # =========================== PARCELAMENTOS ========================

    _add_heading(doc, "PARCELAMENTOS")
    
    # SISPAR - Nova estrutura com parcelamentos
    if "receita_federal" in dados and dados["receita_federal"]:
        receita = dados["receita_federal"]
        sispar = receita.get("sispar", {})
        
        if sispar.get("tem_sispar"):
            parcelamentos = sispar.get("parcelamentos", [])
            
            for idx, parc in enumerate(parcelamentos):
                titulo = f"Parcelamento SISPAR {idx + 1 if len(parcelamentos) > 1 else ''}"
                _add_heading(doc, titulo)
                
                # Informações básicas extraídas do PDF
                conta = parc.get("conta")
                tipo = parc.get("tipo")
                if conta:
                    if tipo:
                        _add_paragrafo(doc, f"Conta: {conta} {tipo}")
                    else:
                        _add_paragrafo(doc, f"Conta: {conta}")
                
                modalidade = parc.get("modalidade")
                if modalidade:
                    _add_paragrafo(doc, f"Modalidade: {modalidade}")
                
                regime = parc.get("regime")
                if regime:
                    _add_paragrafo(doc, f"Regime: {regime}")
                
                limite = parc.get("limite_maximo_meses")
                if limite:
                    _add_paragrafo(doc, f"Limite máximo: ATÉ {limite} MESES")
                
                negociado = parc.get("negociado_no_sispar")
                if negociado is not None:
                    _add_paragrafo(doc, f"Negociado no SISPAR: {'SIM' if negociado else 'NÃO'}")
                
                exigibilidade = parc.get("exigibilidade_suspensa")
                if exigibilidade is not None:
                    _add_paragrafo(doc, f"Exigibilidade suspensa: {'SIM' if exigibilidade else 'NÃO'}")
                
                doc.add_paragraph("")
                
                # Informações preenchidas manualmente (se houver)
                qtd_parcelas = parc.get("quantidade_parcelas")
                valor_total = parc.get("valor_total_parcelado")
                valor_parcela = parc.get("valor_parcela")
                competencias = parc.get("competencias", [])
                
                if qtd_parcelas or valor_total or valor_parcela or competencias:
                    _add_paragrafo(doc, "Informações preenchidas manualmente:")
                    
                    if qtd_parcelas:
                        _add_paragrafo(doc, f"Quantidade de parcelas: {qtd_parcelas}")
                    if valor_total:
                        _add_paragrafo(doc, f"Valor total parcelado: {valor_total}")
                    if valor_parcela:
                        _add_paragrafo(doc, f"Valor da parcela: {valor_parcela}")
                    if competencias:
                        comps_str = ", ".join(competencias)
                        _add_paragrafo(doc, f"Competências: {comps_str}")
                    
                    _add_paragrafo(doc, "Status: INFORMADO PELO USUÁRIO")
                else:
                    # Observação de necessidade de consulta manual
                    observacao = parc.get("observacao", "O relatório da Receita Federal não informa quantidade de parcelas, valores ou competências; é necessária consulta manual ao PGFN/SISPAR.")
                    _add_paragrafo(doc, f"Observação: {observacao}")
                    _add_paragrafo(doc, "Status: NECESSITA CONSULTA MANUAL")
                
                doc.add_paragraph("")
    
    # Parcelamentos manuais
    parcelamentos_rows = dados.get("parcelamentos_rows") or []
    if parcelamentos_rows:
        _add_heading(doc, "Outros Parcelamentos")
        _add_table(
            doc,
            [
                "Parcelamento",
                "Valor aproximado das parcelas",
                "Vencimento",
                "Qtd de parcelas",
                "Parcela atual",
            ],
            parcelamentos_rows,
        )
    elif not ("receita_federal" in dados and dados["receita_federal"] and dados["receita_federal"].get("sispar", {}).get("tem_sispar")):
        _add_paragrafo(doc, "Não há parcelamentos informados.")
    
    # Parcelamentos ativos (manual)
    manual_parc = dados.get("parcelamentos_ativos", {}).get("texto_manual", "").strip()
    if manual_parc:
        _add_paragrafo(doc, "Parcelamentos ativos (manual):")
        for linha in manual_parc.split("\n"):
            if linha.strip():
                _add_paragrafo(doc, linha.strip())
    else:
        _add_paragrafo(doc, "Parcelamentos ativos (manual): (não informado)")
    doc.add_paragraph("")
    
    # ============================= CONCLUSÃO ==========================

    _add_heading(doc, "CONCLUSÃO")
    for linha in dados.get("bloco_conclusao", "").split("\n"):
        if linha.strip():
            _add_paragrafo(doc, linha.strip())
    doc.add_paragraph("")

    _add_paragrafo(doc, "Eikon Soluções Ltda CNPJ: 09.502.539/0001-13")
    doc.add_paragraph("")
    _add_paragrafo(doc, "Atenciosamente,")
    doc.add_paragraph("")
    _add_paragrafo(doc, dados.get("responsavel_nome", ""))
    _add_paragrafo(doc, dados.get("responsavel_cargo", ""))
    _add_paragrafo(doc, f"e-mail: {dados.get('responsavel_email', '')}")

    # ========================= SAÍDA EM MEMÓRIA =======================

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
